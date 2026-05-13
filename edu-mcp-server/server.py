#!/usr/bin/env python3
"""
EDU MCP Server — Expone funcionalidades core de EDU como tools MCP (stdio).

Uso:
  python edu-mcp-server/server.py

Configurar en .vscode/mcp.json:
  {
    "servers": {
      "edu": {
        "type": "stdio",
        "command": "python",
        "args": ["edu-mcp-server/server.py"]
      }
    }
  }

Tools expuestos:
  - edu.search_memory: Buscar en la memoria colectiva (SQLite FTS5)
  - edu.validate_plan: Validar un plan de filminas
  - edu.get_slide_template: Obtener template de slide por tipo
  - edu.search_knowledge: Buscar en ChromaDB knowledge base
"""

from __future__ import annotations

import json
import os
import sqlite3
import subprocess
import sys
from pathlib import Path

try:
    from mcp.server import Server
    from mcp.server.stdio import stdio_server
    from mcp.types import Tool, TextContent
except ImportError:
    print(
        "Error: mcp-sdk no instalado. Ejecutar:\n"
        "  pip install mcp\n"
        "Luego reiniciar el servidor.",
        file=sys.stderr,
    )
    sys.exit(1)


def _find_project_root() -> Path:
    """Busca la raíz del proyecto subiendo desde el directorio actual."""
    current = Path(__file__).resolve().parent.parent
    for candidate in [current, current.parent]:
        if (candidate / "salida" / "edu-standalone").exists():
            return candidate
    return current


PROJECT_ROOT = _find_project_root()
EDU_ROOT = PROJECT_ROOT / "salida" / "edu-standalone"
MEMORY_DB = PROJECT_ROOT / "_edu-memory" / "memory.db"

server = Server("edu-mcp-server")


@server.list_tools()
async def list_tools() -> list[Tool]:
    return [
        Tool(
            name="edu.search_memory",
            description="Buscar en la memoria colectiva EDU (SQLite FTS5). Retorna entries que matchean la query.",
            inputSchema={
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "Texto a buscar"},
                    "limit": {"type": "integer", "default": 10},
                },
                "required": ["query"],
            },
        ),
        Tool(
            name="edu.validate_plan",
            description="Validar un plan de filminas JSON contra el schema EDU.",
            inputSchema={
                "type": "object",
                "properties": {
                    "topic": {"type": "string", "description": "ID del tema (ej: 01-intro)"},
                    "course": {"type": "string", "description": "ID del curso (ej: leng-2026)"},
                },
                "required": ["topic", "course"],
            },
        ),
        Tool(
            name="edu.get_slide_template",
            description="Obtener el template de slide por tipo de filmina.",
            inputSchema={
                "type": "object",
                "properties": {
                    "slide_type": {
                        "type": "string",
                        "enum": [
                            "titulo", "concepto-abstracto", "codigo",
                            "diagrama", "socratica", "cierre",
                        ],
                    },
                },
                "required": ["slide_type"],
            },
        ),
        Tool(
            name="edu.search_knowledge",
            description="Buscar en la knowledge base ChromaDB de EDU (referencias académicas + herramientas).",
            inputSchema={
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "Texto a buscar"},
                    "n_results": {"type": "integer", "default": 5},
                    "doc_type": {
                        "type": "string",
                        "enum": ["reference", "tool", "all"],
                        "default": "all",
                    },
                },
                "required": ["query"],
            },
        ),
        Tool(
            name="edu.generar_oficio_docx",
            description=(
                "Genera un archivo .docx de oficio judicial a partir de la salida estructurada "
                "de ChatGPT. Espera campos: numero, destinatario, asunto, cuerpo, firma. "
                "Guarda el archivo en la carpeta del tema activo."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "numero":       {"type": "string", "description": "Número de oficio"},
                    "destinatario": {"type": "string", "description": "Destinatario del oficio"},
                    "asunto":       {"type": "string", "description": "Asunto del oficio"},
                    "cuerpo":       {"type": "string", "description": "Cuerpo del oficio"},
                    "firma":        {"type": "string", "description": "Firma / firmante"},
                    "ruta_salida":  {
                        "type": "string",
                        "description": "Ruta relativa al .docx generado (opcional). Default: oficio_generado.docx en directorio de trabajo.",
                    },
                },
                "required": ["numero", "destinatario", "asunto", "cuerpo", "firma"],
            },
        ),
        Tool(
            name="edu.anonimizar_documento",
            description=(
                "Anonimiza un texto usando Microsoft Presidio. Detecta y reemplaza PII "
                "(nombres, organizaciones, ubicaciones, fechas, DNI, CUIT, teléfonos) con "
                "marcadores {{TIPO_N}}. Guarda el mapping real→marcador en un archivo JSON local "
                "para que edu.generar_sentencia_docx pueda revertirlo al generar el .docx."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "texto": {
                        "type": "string",
                        "description": "Texto original con datos reales a anonimizar.",
                    },
                    "ruta_mapping": {
                        "type": "string",
                        "description": "Ruta donde guardar el mapping JSON (opcional). Default: anon_mapping.json en directorio de trabajo.",
                    },
                    "idioma": {
                        "type": "string",
                        "default": "es",
                        "description": "Idioma del texto. Default: 'es' (español).",
                    },
                },
                "required": ["texto"],
            },
        ),
        Tool(
            name="edu.generar_sentencia_docx",
            description=(
                "Genera un .docx de sentencia judicial en lenguaje claro a partir de la salida "
                "estructurada de ChatGPT (prompt-sentencia-lenguaje-claro.md). "
                "Campos: expediente, tribunal, fecha, partes, objeto, hechos, "
                "fundamento_juridico, resolutivo, costas, firma."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "expediente":          {"type": "string"},
                    "tribunal":            {"type": "string"},
                    "fecha":               {"type": "string"},
                    "partes":              {"type": "string"},
                    "objeto":              {"type": "string"},
                    "hechos":              {"type": "string"},
                    "fundamento_juridico": {"type": "string"},
                    "resolutivo":          {"type": "string"},
                    "costas":              {"type": "string"},
                    "firma":               {"type": "string"},
                    "ruta_salida":         {
                        "type": "string",
                        "description": "Ruta relativa al .docx (opcional). Default: sentencia_generada.docx",
                    },
                    "ruta_mapping": {
                        "type": "string",
                        "description": "Ruta al mapping JSON generado por edu.anonimizar_documento. Si se provee, los marcadores {{TIPO_N}} se reemplazan con los valores reales antes de generar el .docx.",
                    },
                },
                "required": ["expediente", "tribunal", "fecha", "partes", "objeto",
                             "hechos", "fundamento_juridico", "resolutivo", "costas", "firma"],
            },
        ),
    ]


@server.call_tool()
async def call_tool(name: str, arguments: dict) -> list[TextContent]:
    if name == "edu.search_memory":
        return _search_memory(arguments["query"], arguments.get("limit", 10))
    elif name == "edu.validate_plan":
        return _validate_plan(arguments["topic"], arguments["course"])
    elif name == "edu.get_slide_template":
        return _get_slide_template(arguments["slide_type"])
    elif name == "edu.anonimizar_documento":
        return _anonimizar_documento(arguments)
    elif name == "edu.generar_oficio_docx":
        return _generar_oficio_docx(arguments)
    elif name == "edu.generar_sentencia_docx":
        return _generar_sentencia_docx(arguments)
    elif name == "edu.search_knowledge":
        return _search_knowledge(
            arguments["query"],
            arguments.get("n_results", 5),
            arguments.get("doc_type", "all"),
        )
    else:
        return [TextContent(type="text", text=f"Tool desconocido: {name}")]


def _search_memory(query: str, limit: int) -> list[TextContent]:
    if not MEMORY_DB.exists():
        return [TextContent(type="text", text="memory.db no encontrado.")]
    conn = sqlite3.connect(str(MEMORY_DB))
    try:
        cursor = conn.execute(
            "SELECT category, content, timestamp FROM memory_entries "
            "WHERE content MATCH ? ORDER BY rank LIMIT ?",
            (query, limit),
        )
        rows = cursor.fetchall()
        if not rows:
            return [TextContent(type="text", text=f"Sin resultados para: {query}")]
        results = []
        for cat, content, ts in rows:
            results.append(f"[{cat}] ({ts})\n{content[:300]}")
        return [TextContent(type="text", text="\n---\n".join(results))]
    except sqlite3.OperationalError:
        return [TextContent(type="text", text="Tabla memory_entries no existe aún.")]
    finally:
        conn.close()


def _validate_plan(topic: str, course: str) -> list[TextContent]:
    script = EDU_ROOT / "scripts" / "validate_plan.py"
    if not script.exists():
        return [TextContent(type="text", text="validate_plan.py no encontrado.")]
    result = subprocess.run(
        [sys.executable, str(script), "--topic", topic, "--course", course],
        capture_output=True,
        text=True,
        cwd=str(PROJECT_ROOT),
        timeout=60,
    )
    output = result.stdout or result.stderr or "Sin output"
    return [TextContent(type="text", text=output[:2000])]


def _get_slide_template(slide_type: str) -> list[TextContent]:
    templates = {
        "titulo": "# {title}\n\n**Subtítulo:** {subtitle}\n**Imagen:** {image_description}",
        "concepto-abstracto": "# {title}\n\n{body}\n\n**Imagen:** {image_description}\n**Alt text:** {alt_text}",
        "codigo": "# {title}\n\n```{language}\n{code}\n```\n\n**Notas:** {notes}",
        "diagrama": "# {title}\n\n**Diagrama:** {diagram_description}\n**Nodos:** max 7 (Miller)\n**Flechas:** con labels",
        "socratica": "# {title} ← pregunta\n\n1. {opcion_1}\n2. {opcion_2}\n3. {opcion_3}\n\n**Pausa sugerida:** 30 seg",
        "cierre": "# {title}\n\n**Resumen:** {key_points}\n**Próxima clase:** {next_topic}",
    }
    template = templates.get(slide_type, "Tipo de slide no reconocido.")
    return [TextContent(type="text", text=template)]


def _search_knowledge(query: str, n_results: int, doc_type: str) -> list[TextContent]:
    script = EDU_ROOT / "scripts" / "knowledge_base.py"
    if not script.exists():
        return [TextContent(type="text", text="knowledge_base.py no encontrado.")]
    args = [sys.executable, str(script), "search", query, "--n", str(n_results)]
    if doc_type != "all":
        args.extend(["--type", doc_type])
    result = subprocess.run(
        args, capture_output=True, text=True, cwd=str(PROJECT_ROOT), timeout=60
    )
    output = result.stdout or result.stderr or "Sin resultados"
    return [TextContent(type="text", text=output[:3000])]


def _anonimizar_documento(args: dict) -> list[TextContent]:
    try:
        from presidio_analyzer import AnalyzerEngine
        from presidio_anonymizer import AnonymizerEngine
        from presidio_anonymizer.entities import OperatorConfig
    except ImportError:
        return [TextContent(type="text", text=(
            "❌ Microsoft Presidio no instalado. Ejecutar:\n"
            "  pip install presidio-analyzer presidio-anonymizer\n"
            "  python -m spacy download es_core_news_lg"
        ))]

    texto = args["texto"]
    idioma = args.get("idioma", "es")
    ruta_mapping = Path(args.get("ruta_mapping") or "anon_mapping.json")
    if not ruta_mapping.is_absolute():
        ruta_mapping = PROJECT_ROOT / ruta_mapping

    analyzer = AnalyzerEngine()
    anonymizer = AnonymizerEngine()

    # Tipos de PII a detectar (relevantes para documentos judiciales)
    entidades = [
        "PERSON", "ORG", "LOCATION", "DATE_TIME",
        "ID", "PHONE_NUMBER", "EMAIL_ADDRESS", "NRP",
    ]

    resultados = analyzer.analyze(text=texto, language=idioma, entities=entidades)

    # Construir mapping: marcador → valor real
    mapping: dict[str, str] = {}
    contadores: dict[str, int] = {}
    operadores: dict[str, OperatorConfig] = {}

    # Ordenar por posición para asignar números consistentes
    resultados_ordenados = sorted(resultados, key=lambda r: r.start)

    for res in resultados_ordenados:
        tipo = res.entity_type
        valor_real = texto[res.start:res.end]
        # Reusar marcador si el mismo valor ya fue visto
        marcador_existente = next(
            (m for m, v in mapping.items() if v == valor_real), None
        )
        if not marcador_existente:
            contadores[tipo] = contadores.get(tipo, 0) + 1
            marcador = f"{{{{{tipo}_{contadores[tipo]}}}}}"
            mapping[marcador] = valor_real

    # Anonimizar: reemplazar cada entidad con su marcador
    texto_anonimizado = texto
    # Procesar de atrás hacia adelante para no desplazar índices
    for res in sorted(resultados_ordenados, key=lambda r: r.start, reverse=True):
        tipo = res.entity_type
        valor_real = texto[res.start:res.end]
        marcador = next(
            (m for m, v in mapping.items() if v == valor_real), None
        )
        if marcador:
            texto_anonimizado = (
                texto_anonimizado[:res.start] + marcador + texto_anonimizado[res.end:]
            )

    # Guardar mapping localmente
    import json
    ruta_mapping.parent.mkdir(parents=True, exist_ok=True)
    ruta_mapping.write_text(
        json.dumps({"mapping": mapping, "idioma": idioma}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    resultado = (
        f"✅ Documento anonimizado. Mapping guardado en: {ruta_mapping}\n"
        f"   Entidades detectadas: {len(mapping)}\n\n"
        f"--- TEXTO ANONIMIZADO ---\n{texto_anonimizado}"
    )
    return [TextContent(type="text", text=resultado)]


def _aplicar_mapping(texto: str, ruta_mapping: str) -> str:
    """Reemplaza marcadores {{TIPO_N}} con los valores reales del mapping."""
    import json
    ruta = Path(ruta_mapping)
    if not ruta.is_absolute():
        ruta = PROJECT_ROOT / ruta
    if not ruta.exists():
        return texto  # Si no hay mapping, devolver sin cambios
    data = json.loads(ruta.read_text(encoding="utf-8"))
    mapping = data.get("mapping", {})
    for marcador, valor_real in mapping.items():
        texto = texto.replace(marcador, valor_real)
    return texto


def _generar_sentencia_docx(args: dict) -> list[TextContent]:
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return [TextContent(type="text", text="❌ python-docx no instalado. Ejecutar: pip install python-docx")]

    # Si hay mapping de anonimización, revertir marcadores → valores reales
    ruta_mapping = args.get("ruta_mapping")
    def deanon(texto: str) -> str:
        return _aplicar_mapping(texto, ruta_mapping) if ruta_mapping else texto

    # Revertir en todos los campos de texto
    campos_texto = ["expediente", "tribunal", "fecha", "partes", "objeto",
                    "hechos", "fundamento_juridico", "resolutivo", "costas", "firma"]
    args = {k: (deanon(v) if k in campos_texto and isinstance(v, str) else v)
            for k, v in args.items()}

    ruta_salida = Path(args.get("ruta_salida") or "sentencia_generada.docx")
    if not ruta_salida.is_absolute():
        ruta_salida = PROJECT_ROOT / ruta_salida

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.2)
        section.bottom_margin = Inches(1.2)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.2)

    def heading(text: str, level: int = 1):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11 if level == 1 else 10)
        return p

    def field(label: str, value: str):
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)
        return p

    # Encabezado institucional
    enc = doc.add_paragraph()
    enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = enc.add_run("PODER JUDICIAL DE LA NACIÓN")
    r.bold = True
    r.font.size = Pt(13)

    trib = doc.add_paragraph()
    trib.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trib.add_run(args["tribunal"]).font.size = Pt(11)

    doc.add_paragraph()

    # Metadatos
    field("Expediente", args["expediente"])
    field("Fecha", args["fecha"])
    field("Partes", args["partes"])
    field("Objeto", args["objeto"])

    doc.add_paragraph()

    # Resolutivo primero (lenguaje claro: decisión al inicio)
    heading("RESOLUTIVO")
    resolutivo_p = doc.add_paragraph(args["resolutivo"])
    resolutivo_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # Hechos
    heading("HECHOS PROBADOS")
    hechos_p = doc.add_paragraph(args["hechos"])
    hechos_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # Fundamento jurídico
    heading("FUNDAMENTO JURÍDICO")
    fund_p = doc.add_paragraph(args["fundamento_juridico"])
    fund_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # Costas
    field("Costas", args["costas"])

    doc.add_paragraph()
    doc.add_paragraph()

    # Firma
    firma_p = doc.add_paragraph()
    firma_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma_p.add_run(args["firma"]).bold = True

    doc.save(str(ruta_salida))
    return [TextContent(type="text", text=f"✅ Sentencia generada: {ruta_salida}")]


def _generar_oficio_docx(args: dict) -> list[TextContent]:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return [TextContent(type="text", text="❌ python-docx no instalado. Ejecutar: pip install python-docx")]

    ruta_salida = Path(args.get("ruta_salida") or "oficio_generado.docx")
    if not ruta_salida.is_absolute():
        ruta_salida = PROJECT_ROOT / ruta_salida

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.2)
        section.bottom_margin = Inches(1.2)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.2)

    enc = doc.add_paragraph()
    enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = enc.add_run("PODER JUDICIAL DE LA NACIÓN")
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    nro = doc.add_paragraph()
    nro.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    nro.add_run(f"Oficio N°: {args['numero']}").bold = True

    doc.add_paragraph()

    dest = doc.add_paragraph()
    dest.add_run("A: ").bold = True
    dest.add_run(args["destinatario"])

    asunto_p = doc.add_paragraph()
    asunto_p.add_run("Asunto: ").bold = True
    asunto_p.add_run(args["asunto"])

    doc.add_paragraph()

    cuerpo_p = doc.add_paragraph(args["cuerpo"])
    cuerpo_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    doc.add_paragraph()

    firma_p = doc.add_paragraph()
    firma_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma_p.add_run(args["firma"]).bold = True

    doc.save(str(ruta_salida))
    return [TextContent(type="text", text=f"✅ Oficio generado: {ruta_salida}")]


async def main():
    async with stdio_server() as (read_stream, write_stream):
        await server.run(read_stream, write_stream, server.create_initialization_options())


if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
