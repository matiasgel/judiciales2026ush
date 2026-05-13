"""
oficio_desde_template.py
Demo de clase — Encuentro 1: "Entender para no equivocarse"

USO:
    1. Copiá la respuesta de ChatGPT en un archivo de texto (respuesta.txt)
       o pegala cuando el script la pida por stdin.
    2. Ejecutá: python oficio_desde_template.py [respuesta.txt]
    3. Se genera: oficio_generado.docx

FORMATO ESPERADO DE LA RESPUESTA DE CHATGPT:
    El prompt que le diste a ChatGPT debe incluir esta plantilla de salida:

    OFICIO N°: {{numero}}
    Destinatario: {{destinatario}}
    Asunto: {{asunto}}
    Cuerpo: {{cuerpo}}
    Firma: {{firma}}

    El script parsea esos campos y los vuelca al .docx.

DEPENDENCIA:
    pip install python-docx
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


CAMPOS = ["numero", "destinatario", "asunto", "cuerpo", "firma"]

ETIQUETAS = {
    "numero":       "OFICIO N°:",
    "destinatario": "Destinatario:",
    "asunto":       "Asunto:",
    "cuerpo":       "Cuerpo:",
    "firma":        "Firma:",
}


def parsear_respuesta(texto: str) -> dict:
    """Extrae los campos de la respuesta estructurada de ChatGPT."""
    valores = {c: "" for c in CAMPOS}
    lineas = texto.splitlines()

    campo_actual = None
    buffer = []

    for linea in lineas:
        match = None
        for campo, etiqueta in ETIQUETAS.items():
            if linea.strip().upper().startswith(etiqueta.upper()):
                if campo_actual and buffer:
                    valores[campo_actual] = " ".join(buffer).strip()
                campo_actual = campo
                resto = linea.strip()[len(etiqueta):].strip()
                buffer = [resto] if resto else []
                match = True
                break
        if not match and campo_actual:
            buffer.append(linea.strip())

    if campo_actual and buffer:
        valores[campo_actual] = " ".join(buffer).strip()

    return valores


def generar_docx(campos: dict, ruta_salida: str = "oficio_generado.docx"):
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1.2)
        section.bottom_margin = Inches(1.2)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.2)

    # Encabezado institucional
    enc = doc.add_paragraph()
    enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = enc.add_run("PODER JUDICIAL DE LA NACIÓN")
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()  # espacio

    # Número de oficio
    nro = doc.add_paragraph()
    nro.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    nro.add_run(f"Oficio N°: {campos['numero']}").bold = True

    doc.add_paragraph()

    # Destinatario
    dest = doc.add_paragraph()
    dest.add_run("A: ").bold = True
    dest.add_run(campos["destinatario"])

    # Asunto
    asunto = doc.add_paragraph()
    asunto.add_run("Asunto: ").bold = True
    asunto.add_run(campos["asunto"])

    doc.add_paragraph()

    # Cuerpo
    cuerpo_p = doc.add_paragraph(campos["cuerpo"])
    cuerpo_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    doc.add_paragraph()

    # Firma
    firma_p = doc.add_paragraph()
    firma_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma_p.add_run(campos["firma"]).bold = True

    doc.save(ruta_salida)
    print(f"✅ Documento generado: {ruta_salida}")


def main():
    if len(sys.argv) > 1:
        archivo = Path(sys.argv[1])
        if not archivo.exists():
            print(f"❌ No se encontró el archivo: {archivo}")
            sys.exit(1)
        texto = archivo.read_text(encoding="utf-8")
    else:
        print("Pegá la respuesta de ChatGPT y presioná Enter dos veces + Ctrl+Z (Windows) o Ctrl+D (Unix):")
        lineas = []
        try:
            for linea in sys.stdin:
                lineas.append(linea)
        except KeyboardInterrupt:
            pass
        texto = "".join(lineas)

    if not texto.strip():
        print("❌ No se recibió texto. Asegurate de pegar la respuesta de ChatGPT.")
        sys.exit(1)

    campos = parsear_respuesta(texto)

    faltantes = [c for c in CAMPOS if not campos[c]]
    if faltantes:
        print(f"⚠️  Campos no encontrados en la respuesta: {', '.join(faltantes)}")
        print("   Verificá que el prompt incluyó la plantilla de salida correctamente.")

    generar_docx(campos)


if __name__ == "__main__":
    main()
